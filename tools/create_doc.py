"""
Creates a professional Word document (.docx) explaining the Blog Auto-Poster system.
Can be uploaded to Google Docs directly.
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import os

doc = Document()

# ============================================================
# STYLES
# ============================================================
style = doc.styles["Normal"]
font = style.font
font.name = "Calibri"
font.size = Pt(11)

# ============================================================
# TITLE PAGE
# ============================================================
for _ in range(6):
    doc.add_paragraph("")

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run("BLOG AUTO-POSTER SYSTEM")
run.bold = True
run.font.size = Pt(28)
run.font.color.rgb = RGBColor(0, 102, 204)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run("Complete Documentation & Setup Guide")
run.font.size = Pt(16)
run.font.color.rgb = RGBColor(100, 100, 100)

doc.add_paragraph("")

desc = doc.add_paragraph()
desc.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = desc.add_run(
    "Automated blog posting system for SEO backlink building.\n"
    "Generates unique articles and publishes to 6 free blog platforms."
)
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(80, 80, 80)

doc.add_paragraph("")
doc.add_paragraph("")

info = doc.add_paragraph()
info.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = info.add_run("Version 1.0  |  2026")
run.font.size = Pt(11)
run.font.color.rgb = RGBColor(150, 150, 150)

doc.add_page_break()

# ============================================================
# TABLE OF CONTENTS
# ============================================================
doc.add_heading("Table of Contents", level=1)
toc_items = [
    "1. What Is This System?",
    "2. How It Works (Flow Diagram)",
    "3. Supported Platforms (6 Platforms)",
    "4. Quick Start Guide",
    "5. Platform Setup (Step by Step)",
    "6. Content Generation Methods",
    "7. Backlink Injection Strategy",
    "8. Anti-Spam Protection",
    "9. All Commands & Usage",
    "10. File Structure",
    "11. Configuration Reference",
    "12. Scheduling & Automation",
    "13. Tracking & Reporting",
    "14. Troubleshooting",
    "15. SEO Best Practices & Tips",
]
for item in toc_items:
    p = doc.add_paragraph(item)
    p.paragraph_format.space_after = Pt(4)

doc.add_page_break()

# ============================================================
# SECTION 1: WHAT IS THIS
# ============================================================
doc.add_heading("1. What Is This System?", level=1)

doc.add_heading("The Problem", level=2)
doc.add_paragraph(
    "Google ranks websites higher when other websites link back to them (backlinks). "
    "Manually creating blog posts on 6 different platforms and writing unique articles "
    "with backlinks is extremely time-consuming. A single round of posting across all "
    "platforms can take 3-4 hours of manual work."
)

doc.add_heading("The Solution", level=2)
doc.add_paragraph(
    "The Blog Auto-Poster is an automated system that handles everything for you:"
)

bullets = [
    "Generates unique blog articles using templates (and optionally AI)",
    "Injects your website backlinks naturally inside each article",
    "Publishes automatically to 6 free blog platforms via their APIs",
    "Tracks every published post with URLs in CSV/Excel reports",
    "Schedules daily auto-posting so you don't have to do anything",
    "Uses anti-spam measures to keep your posts looking natural",
]
for b in bullets:
    doc.add_paragraph(b, style="List Bullet")

doc.add_heading("Who Is This For?", level=2)
doc.add_paragraph(
    "This tool is designed for business owners, digital marketers, and SEO professionals "
    "who want to build backlinks to their website through blog posting on high-authority "
    "platforms. It's especially useful for local businesses wanting to improve their "
    "Google search rankings."
)

doc.add_page_break()

# ============================================================
# SECTION 2: HOW IT WORKS
# ============================================================
doc.add_heading("2. How It Works", level=1)

doc.add_heading("Complete Flow", level=2)

steps = [
    ("Step 1: Configure", "You set your website URL, brand name, and target keywords in blog_config.py"),
    ("Step 2: Generate Topic", "System auto-generates a unique blog topic from your keywords\nExample: 'Top 10 Digital Marketing Strategies for 2026'"),
    ("Step 3: Generate Article", "System creates a 500-1000 word article using templates.\nContent spinning makes each version unique."),
    ("Step 4: Inject Backlinks", "System naturally places 2 backlinks to your website inside the article.\nAlso adds an author bio with your website link."),
    ("Step 5: Publish", "System posts the article to the selected platform via API.\nSupports: Telegraph, WordPress, Blogger, Tumblr, Medium, Hashnode"),
    ("Step 6: Track", "System logs the post URL, platform, date, status in CSV.\nYou can generate an Excel report anytime."),
    ("Step 7: Wait & Repeat", "System waits 5-15 minutes (random delay), then posts to the next platform.\nThis makes the posting look natural and human-like."),
]

table = doc.add_table(rows=len(steps) + 1, cols=2)
table.style = "Light Grid Accent 1"
table.alignment = WD_TABLE_ALIGNMENT.CENTER

# Header row
hdr = table.rows[0]
hdr.cells[0].text = "Step"
hdr.cells[1].text = "What Happens"
for cell in hdr.cells:
    cell.paragraphs[0].runs[0].bold = True

for i, (step, desc) in enumerate(steps):
    row = table.rows[i + 1]
    row.cells[0].text = step
    row.cells[1].text = desc

doc.add_paragraph("")

doc.add_page_break()

# ============================================================
# SECTION 3: PLATFORMS
# ============================================================
doc.add_heading("3. Supported Platforms", level=1)

doc.add_paragraph(
    "The system supports 6 free blog platforms. Each platform has different strengths "
    "for SEO. You can enable/disable any platform in the configuration."
)

platforms = [
    ("Telegraph (telegra.ph)", "EASIEST", "Medium", "No signup needed, no API key. Token auto-generated. "
     "Anonymous blogging platform by Telegram. Full HTML support with dofollow links."),
    ("WordPress.com", "BEST FOR SEO", "93 (Very High)", "Free blog at yourblog.wordpress.com. "
     "Highest domain authority among free platforms. Google trusts WordPress.com backlinks. "
     "Requires OAuth2 token from developer.wordpress.com."),
    ("Blogger (Google)", "HIGH DA", "89 (Very High)", "Free blog at yourblog.blogspot.com. "
     "Owned by Google, excellent for SEO. Requires Google Cloud API key and Blog ID."),
    ("Tumblr", "EASY API", "86 (High)", "Microblogging platform with good DA. "
     "DoFollow links in posts. Requires OAuth1 credentials from tumblr.com/oauth/apps."),
    ("Medium", "HIGH AUTHORITY", "95 (Very High)", "Popular blogging platform with massive audience. "
     "Very high DA but links may be nofollow. Simple bearer token auth. API is limited."),
    ("Hashnode", "DEVELOPER BLOG", "72 (Good)", "Developer-focused blogging platform. "
     "DoFollow links, free custom domain. Uses GraphQL API with personal access token."),
]

table = doc.add_table(rows=len(platforms) + 1, cols=4)
table.style = "Light Grid Accent 1"
table.alignment = WD_TABLE_ALIGNMENT.CENTER

hdr = table.rows[0]
headers = ["Platform", "Difficulty", "Domain Authority", "Details"]
for j, h in enumerate(headers):
    hdr.cells[j].text = h
    hdr.cells[j].paragraphs[0].runs[0].bold = True

for i, (name, diff, da, details) in enumerate(platforms):
    row = table.rows[i + 1]
    row.cells[0].text = name
    row.cells[1].text = diff
    row.cells[2].text = da
    row.cells[3].text = details

doc.add_paragraph("")
doc.add_paragraph(
    "Recommendation: Start with Telegraph (zero setup), then add WordPress.com and "
    "Blogger for maximum SEO impact. Add other platforms gradually."
)

doc.add_page_break()

# ============================================================
# SECTION 4: QUICK START
# ============================================================
doc.add_heading("4. Quick Start Guide (5 Minutes)", level=1)

doc.add_heading("Step 1: Edit Configuration", level=2)
doc.add_paragraph("Open blog_poster/blog_config.py and change these values:")
doc.add_paragraph(
    'TARGET_URL = "https://yourwebsite.com"        # Your website\n'
    'TARGET_BRAND = "Your Brand Name"               # Your brand\n'
    'TARGET_KEYWORDS = ["keyword1", "keyword2", ...]  # Your SEO keywords',
    style="No Spacing",
)

doc.add_heading("Step 2: Enable Platform", level=2)
doc.add_paragraph(
    "Telegraph is already enabled by default (no signup needed). "
    "For other platforms, set 'enabled': True and add credentials."
)

doc.add_heading("Step 3: Run the System", level=2)
commands = [
    ("python -m blog_poster.blog_poster --platform telegraph", "Post to Telegraph"),
    ("python -m blog_poster.blog_poster --all-platforms", "Post to all enabled platforms"),
    ("python -m blog_poster.blog_poster --all-platforms --count 3", "Post 3 articles"),
    ("python -m blog_poster.blog_poster --schedule", "Schedule daily auto-posting"),
    ("python -m blog_poster.blog_poster --report", "Generate Excel report"),
]

table = doc.add_table(rows=len(commands) + 1, cols=2)
table.style = "Light Grid Accent 1"
hdr = table.rows[0]
hdr.cells[0].text = "Command"
hdr.cells[1].text = "Description"
for cell in hdr.cells:
    cell.paragraphs[0].runs[0].bold = True
for i, (cmd, desc) in enumerate(commands):
    table.rows[i + 1].cells[0].text = cmd
    table.rows[i + 1].cells[1].text = desc

doc.add_heading("Step 4: Check Results", level=2)
doc.add_paragraph("Your posts are logged in: blog_poster/output/posts_log.csv")
doc.add_paragraph("Generate a formatted report: python -m blog_poster.blog_poster --report")

doc.add_page_break()

# ============================================================
# SECTION 5: PLATFORM SETUP
# ============================================================
doc.add_heading("5. Platform Setup (Step by Step)", level=1)

# Telegraph
doc.add_heading("5.1 Telegraph Setup", level=2)
doc.add_paragraph("No setup needed! Telegraph works out of the box.", style="List Bullet")
doc.add_paragraph("Token is auto-generated on first run and saved to telegraph_token.txt", style="List Bullet")
doc.add_paragraph("Just run: python -m blog_poster.blog_poster --platform telegraph", style="List Bullet")

# WordPress
doc.add_heading("5.2 WordPress.com Setup", level=2)
wp_steps = [
    "Create a free blog at wordpress.com (choose any .wordpress.com subdomain)",
    "Go to developer.wordpress.com/apps/ and click 'Create New Application'",
    "Fill in app name, description, and website URL",
    "Set Redirect URL to: https://example.com (any URL works)",
    "After creation, note your Client ID and Client Secret",
    "Get your Access Token using the OAuth2 flow",
    "Edit blog_config.py: set site_url and access_token",
    "Set 'enabled': True",
]
for step in wp_steps:
    doc.add_paragraph(step, style="List Number")

# Blogger
doc.add_heading("5.3 Blogger (Google) Setup", level=2)
bl_steps = [
    "Go to console.cloud.google.com",
    "Create a new project (or select existing)",
    "Go to 'APIs & Services' > 'Library'",
    "Search for 'Blogger API v3' and enable it",
    "Go to 'Credentials' > 'Create Credentials' > 'API Key'",
    "Copy the API key",
    "Create a blog at blogger.com",
    "Get your Blog ID from the URL (Settings > Blog ID)",
    "Edit blog_config.py: set blog_id and api_key",
    "Set 'enabled': True",
]
for step in bl_steps:
    doc.add_paragraph(step, style="List Number")

# Tumblr
doc.add_heading("5.4 Tumblr Setup", level=2)
tu_steps = [
    "Create a Tumblr account and blog at tumblr.com",
    "Go to tumblr.com/oauth/apps",
    "Click 'Register application'",
    "Fill in app details (any name and description)",
    "After registration, get: Consumer Key, Consumer Secret",
    "Use the OAuth1 flow to get: OAuth Token, OAuth Token Secret",
    "Edit blog_config.py: fill in all 4 credentials + blog_name",
    "Install dependency: pip install requests-oauthlib",
    "Set 'enabled': True",
]
for step in tu_steps:
    doc.add_paragraph(step, style="List Number")

# Medium
doc.add_heading("5.5 Medium Setup", level=2)
me_steps = [
    "Create a Medium account at medium.com",
    "Go to medium.com/me/settings",
    "Scroll down to 'Integration tokens'",
    "Click 'Get integration token' and copy it",
    "Edit blog_config.py: set integration_token",
    "Set 'enabled': True",
]
for step in me_steps:
    doc.add_paragraph(step, style="List Number")

# Hashnode
doc.add_heading("5.6 Hashnode Setup", level=2)
hn_steps = [
    "Create account at hashnode.com",
    "Start a free blog from your dashboard",
    "Go to hashnode.com/settings/developer",
    "Generate a 'Personal Access Token' and copy it",
    "Get your Publication ID from blog settings",
    "Edit blog_config.py: set token and publication_id",
    "Install dependency: pip install markdownify",
    "Set 'enabled': True",
]
for step in hn_steps:
    doc.add_paragraph(step, style="List Number")

doc.add_page_break()

# ============================================================
# SECTION 6: CONTENT GENERATION
# ============================================================
doc.add_heading("6. Content Generation Methods", level=1)

doc.add_heading("Method 1: Templates (Default)", level=2)
doc.add_paragraph(
    "The default method uses pre-built article templates. No AI or internet needed. "
    "Templates are organized by niche and have placeholders that get filled with your keywords."
)
doc.add_paragraph("Available niches:", style="List Bullet")
niches = ["Business & Marketing", "Local Services", "Technology & Digital",
          "How-To Guides", "Health & Wellness", "Education & Training"]
for n in niches:
    doc.add_paragraph(f"  - {n}")

doc.add_paragraph(
    "\nContent Spinning: The system automatically replaces words with synonyms and "
    "shuffles paragraph order to make each article unique. This prevents duplicate "
    "content penalties from search engines."
)

doc.add_heading("Method 2: Ollama (Free Local AI)", level=2)
doc.add_paragraph(
    "Ollama runs AI models locally on your computer - completely free, no API key needed. "
    "It generates high-quality, unique articles."
)
ollama_steps = [
    "Download and install Ollama from ollama.ai",
    "Open terminal and run: ollama pull llama3",
    "In blog_config.py, set: CONTENT_METHOD = 'ollama'",
    "Run the system normally - it will use AI for content",
]
for step in ollama_steps:
    doc.add_paragraph(step, style="List Number")

doc.add_heading("Method 3: OpenAI (Paid)", level=2)
doc.add_paragraph(
    "If you have an OpenAI API key, you can use GPT-3.5 or GPT-4 for premium content. "
    "Set OPENAI_API_KEY and CONTENT_METHOD = 'openai' in blog_config.py."
)

doc.add_page_break()

# ============================================================
# SECTION 7: BACKLINK STRATEGY
# ============================================================
doc.add_heading("7. Backlink Injection Strategy", level=1)

doc.add_paragraph(
    "The system places backlinks naturally inside articles to avoid looking spammy. "
    "Here's exactly how it works:"
)

doc.add_heading("In-Content Links (2 per article)", level=2)
doc.add_paragraph("One link is placed in the first third of the article", style="List Bullet")
doc.add_paragraph("One link is placed in the last third of the article", style="List Bullet")
doc.add_paragraph("Links are wrapped around relevant keyword phrases", style="List Bullet")
doc.add_paragraph("Anchor text is randomly selected from your configured list", style="List Bullet")
doc.add_paragraph(
    '\nExample: "For best results, consider working with '
    '<a href="https://yoursite.com">Your Brand Name</a> for expert guidance."'
)

doc.add_heading("Author Bio Link", level=2)
doc.add_paragraph(
    "Every article ends with an author bio that includes a backlink to your website. "
    "This is a common and accepted practice in blogging."
)
doc.add_paragraph(
    'Example: "This article was contributed by Your Brand Name, '
    'a trusted name in the industry. Visit us for more insights."'
)

doc.add_heading("Natural Link Profile", level=2)
doc.add_paragraph("70% of links are DoFollow (pass SEO value)", style="List Bullet")
doc.add_paragraph("30% of links are NoFollow (looks natural to Google)", style="List Bullet")
doc.add_paragraph("1-2 outbound links to authority sites (Wikipedia, Forbes) are mixed in", style="List Bullet")
doc.add_paragraph("Different anchor texts used for each post (rotation)", style="List Bullet")

doc.add_page_break()

# ============================================================
# SECTION 8: ANTI-SPAM
# ============================================================
doc.add_heading("8. Anti-Spam Protection", level=1)

doc.add_paragraph(
    "The system includes multiple anti-spam measures to ensure your posts look natural "
    "and don't get flagged by platforms or search engines."
)

spam_measures = [
    ("Unique Content Per Platform", "Each platform receives a different version of the article. Content spinning ensures no two posts are identical."),
    ("Random Delays (5-15 min)", "The system waits a random 5-15 minutes between posts. This mimics human posting behavior."),
    ("Anchor Text Rotation", "Different link text is used for each post. Never the same anchor text twice in a row."),
    ("DoFollow/NoFollow Mix", "30% of backlinks are marked nofollow. A 100% dofollow profile looks unnatural."),
    ("Word Count Variation", "Articles range from 500-1200 words randomly. Uniform length looks automated."),
    ("Platform Rotation", "Posts cycle through platforms in round-robin order. Never posts to the same platform consecutively."),
    ("Daily Post Limits", "Maximum posts per day (default: 3) and per platform (default: 1) are enforced."),
    ("Authority Link Mixing", "Links to trusted sites like Wikipedia are added alongside your backlinks."),
]

table = doc.add_table(rows=len(spam_measures) + 1, cols=2)
table.style = "Light Grid Accent 1"
hdr = table.rows[0]
hdr.cells[0].text = "Protection"
hdr.cells[1].text = "How It Works"
for cell in hdr.cells:
    cell.paragraphs[0].runs[0].bold = True
for i, (measure, detail) in enumerate(spam_measures):
    table.rows[i + 1].cells[0].text = measure
    table.rows[i + 1].cells[1].text = detail

doc.add_page_break()

# ============================================================
# SECTION 9: ALL COMMANDS
# ============================================================
doc.add_heading("9. All Commands & Usage", level=1)

all_commands = [
    ("--platform telegraph", "Post to a specific platform"),
    ("--all-platforms", "Post to all enabled platforms"),
    ("--all-platforms --count 3", "Post 3 articles across platforms"),
    ("--topic 'Custom Topic'", "Use a custom blog topic"),
    ("--niche business", "Use templates from a specific niche"),
    ("--schedule", "Schedule daily auto-posting"),
    ("--report", "Generate Excel report of all posts"),
    ("--validate", "Test all platform credentials"),
]

table = doc.add_table(rows=len(all_commands) + 1, cols=2)
table.style = "Light Grid Accent 1"
hdr = table.rows[0]
hdr.cells[0].text = "Command Flag"
hdr.cells[1].text = "What It Does"
for cell in hdr.cells:
    cell.paragraphs[0].runs[0].bold = True
for i, (cmd, desc) in enumerate(all_commands):
    table.rows[i + 1].cells[0].text = cmd
    table.rows[i + 1].cells[1].text = desc

doc.add_paragraph(
    "\nAll commands start with: python -m blog_poster.blog_poster"
)

doc.add_page_break()

# ============================================================
# SECTION 10: FILE STRUCTURE
# ============================================================
doc.add_heading("10. File Structure", level=1)

files = [
    ("blog_config.py", "All settings - website URL, credentials, delays, content method"),
    ("blog_utils.py", "Content generation, text spinning, backlink injection, formatting"),
    ("blog_platforms.py", "6 platform poster classes + PosterFactory"),
    ("blog_poster.py", "Main pipeline - orchestrates everything"),
    ("blog_tracker.py", "CSV/Excel tracking and reporting"),
    ("templates/article_templates.py", "10+ article templates organized by 6 niches"),
    ("output/posts_log.csv", "All posts logged here (auto-created)"),
    ("output/posts_report.xlsx", "Excel report (generated via --report)"),
    ("output/poster.log", "Debug log for troubleshooting"),
    ("telegraph_token.txt", "Auto-generated Telegraph API token"),
    ("BLOG_POSTER_GUIDE.md", "Quick reference guide (Markdown)"),
]

table = doc.add_table(rows=len(files) + 1, cols=2)
table.style = "Light Grid Accent 1"
hdr = table.rows[0]
hdr.cells[0].text = "File"
hdr.cells[1].text = "Purpose"
for cell in hdr.cells:
    cell.paragraphs[0].runs[0].bold = True
for i, (f, p) in enumerate(files):
    table.rows[i + 1].cells[0].text = f
    table.rows[i + 1].cells[1].text = p

doc.add_page_break()

# ============================================================
# SECTION 11: CONFIG REFERENCE
# ============================================================
doc.add_heading("11. Configuration Reference", level=1)

doc.add_paragraph(
    "All settings are in blog_poster/blog_config.py. Here's every setting explained:"
)

doc.add_heading("Target Website Settings", level=2)
configs = [
    ("TARGET_URL", "Your website URL that needs backlinks"),
    ("TARGET_BRAND", "Your brand/company name"),
    ("TARGET_KEYWORDS", "List of SEO keywords related to your business"),
    ("ANCHOR_TEXTS", "List of anchor texts for backlinks (rotated randomly)"),
]
for name, desc in configs:
    p = doc.add_paragraph()
    run = p.add_run(f"{name}: ")
    run.bold = True
    p.add_run(desc)

doc.add_heading("Content Settings", level=2)
configs = [
    ("CONTENT_METHOD", "'template' (default), 'ollama' (free AI), or 'openai' (paid)"),
    ("MIN_WORD_COUNT", "Minimum article length (default: 500 words)"),
    ("MAX_WORD_COUNT", "Maximum article length (default: 1200 words)"),
    ("SPIN_CONTENT", "Enable/disable content spinning (default: True)"),
]
for name, desc in configs:
    p = doc.add_paragraph()
    run = p.add_run(f"{name}: ")
    run.bold = True
    p.add_run(desc)

doc.add_heading("Posting Schedule Settings", level=2)
configs = [
    ("POSTS_PER_DAY", "Maximum total posts per day (default: 3)"),
    ("MAX_POSTS_PER_PLATFORM", "Maximum posts per platform per day (default: 1)"),
    ("MIN_DELAY_SECONDS", "Minimum wait between posts (default: 300 = 5 min)"),
    ("MAX_DELAY_SECONDS", "Maximum wait between posts (default: 900 = 15 min)"),
    ("DAILY_POST_TIME", "When daily auto-posting starts (default: '10:00')"),
]
for name, desc in configs:
    p = doc.add_paragraph()
    run = p.add_run(f"{name}: ")
    run.bold = True
    p.add_run(desc)

doc.add_heading("Backlink Settings", level=2)
configs = [
    ("MAX_BACKLINKS_PER_POST", "Max self-links per article (default: 2)"),
    ("ADD_AUTHOR_BIO", "Add author bio with link at end (default: True)"),
    ("NOFOLLOW_PERCENTAGE", "% of links marked nofollow (default: 30)"),
    ("AUTHORITY_LINKS", "External authority sites to link to (looks natural)"),
]
for name, desc in configs:
    p = doc.add_paragraph()
    run = p.add_run(f"{name}: ")
    run.bold = True
    p.add_run(desc)

doc.add_page_break()

# ============================================================
# SECTION 12: SCHEDULING
# ============================================================
doc.add_heading("12. Scheduling & Automation", level=1)

doc.add_paragraph(
    "The system can run on a daily schedule, automatically posting articles "
    "at your configured time without any manual intervention."
)

doc.add_heading("Start Daily Auto-Posting", level=2)
doc.add_paragraph("python -m blog_poster.blog_poster --schedule", style="No Spacing")
doc.add_paragraph(
    "\nThis will post immediately for the first time, then repeat every day "
    "at the configured time (default: 10:00 AM)."
)

doc.add_heading("Customize Schedule", level=2)
doc.add_paragraph("In blog_config.py:")
doc.add_paragraph(
    'DAILY_POST_TIME = "10:00"   # Change to your preferred time\n'
    'POSTS_PER_DAY = 3           # Total posts per day\n'
    'MAX_POSTS_PER_PLATFORM = 1  # Posts per platform per day',
    style="No Spacing",
)

doc.add_heading("Keep Running in Background", level=2)
doc.add_paragraph(
    "The scheduler needs the terminal to stay open. Options for keeping it running:"
)
doc.add_paragraph("Windows: Use Task Scheduler to run the script daily", style="List Bullet")
doc.add_paragraph("Linux: Use cron job or systemd service", style="List Bullet")
doc.add_paragraph("Cloud: Deploy on a free VPS or cloud function", style="List Bullet")

doc.add_page_break()

# ============================================================
# SECTION 13: TRACKING
# ============================================================
doc.add_heading("13. Tracking & Reporting", level=1)

doc.add_heading("Automatic Post Logging", level=2)
doc.add_paragraph(
    "Every post (success or failure) is automatically logged to posts_log.csv with:"
)
tracking_fields = [
    "Post ID - Unique identifier from the platform",
    "Platform - Which platform it was posted to",
    "Title - The article title",
    "Published URL - Direct link to the live post",
    "Target URL - Your website URL (the backlink destination)",
    "Anchor Text - The clickable text used for the backlink",
    "Tags - Keywords/tags used for the post",
    "Word Count - Article length",
    "Status - 'published' or 'failed'",
    "Error - Error message if post failed",
    "Date Published - When the post was created",
]
for field in tracking_fields:
    doc.add_paragraph(field, style="List Bullet")

doc.add_heading("Excel Report", level=2)
doc.add_paragraph(
    "Generate a formatted Excel report with:\n"
    "python -m blog_poster.blog_poster --report"
)
doc.add_paragraph("The report includes:", style="List Bullet")
doc.add_paragraph("  - Summary sheet (total posts, success rate, platforms used)")
doc.add_paragraph("  - All Posts sheet (every post with clickable URLs)")
doc.add_paragraph("  - Per-platform sheets (filtered by platform)")
doc.add_paragraph("  - Failed Posts sheet (for debugging)")

doc.add_page_break()

# ============================================================
# SECTION 14: TROUBLESHOOTING
# ============================================================
doc.add_heading("14. Troubleshooting", level=1)

issues = [
    ("No platforms enabled", "Edit blog_config.py and set 'enabled': True for at least one platform. Telegraph needs no setup."),
    ("Telegraph connection error", "Check your internet connection. Telegraph API may be blocked in some regions - try a VPN."),
    ("WordPress 401 Unauthorized", "Your access token has expired. Regenerate it from developer.wordpress.com."),
    ("Blogger 403 Forbidden", "Check that Blogger API is enabled in Google Cloud Console and API key is correct."),
    ("Tumblr authentication error", "Verify all 4 OAuth credentials. Install requests-oauthlib if not installed."),
    ("Medium rate limit", "Medium severely limits API usage. Wait 24 hours before trying again."),
    ("Hashnode mutation error", "Verify your token and publication_id. Make sure you have a blog created on Hashnode."),
    ("Content generation empty", "Check that templates exist in templates/article_templates.py and keywords are set in config."),
    ("Import error", "Run: pip install requests-oauthlib markdownify python-docx openpyxl"),
    ("Excel permission error", "Close the Excel file before running the report again."),
]

table = doc.add_table(rows=len(issues) + 1, cols=2)
table.style = "Light Grid Accent 1"
hdr = table.rows[0]
hdr.cells[0].text = "Problem"
hdr.cells[1].text = "Solution"
for cell in hdr.cells:
    cell.paragraphs[0].runs[0].bold = True
for i, (prob, sol) in enumerate(issues):
    table.rows[i + 1].cells[0].text = prob
    table.rows[i + 1].cells[1].text = sol

doc.add_page_break()

# ============================================================
# SECTION 15: SEO TIPS
# ============================================================
doc.add_heading("15. SEO Best Practices & Tips", level=1)

tips = [
    ("Start with Telegraph", "It's the easiest platform with zero setup. Test your entire flow here first before adding other platforms."),
    ("Add platforms gradually", "Don't enable all 6 platforms at once. Add one new platform per week to build a natural backlink profile."),
    ("Use relevant keywords", "Your TARGET_KEYWORDS should match what people search for when looking for your business."),
    ("Vary anchor texts", "Add 5-10 different anchor text options. Over-optimized anchor text is the #1 spam signal for Google."),
    ("Don't overdo it", "2-3 posts per day is plenty. Posting 50 articles a day will get you flagged as spam."),
    ("Be patient", "SEO takes 2-3 months to show results. Keep posting consistently and results will come."),
    ("Monitor rankings", "Use Google Search Console (free) to track how your rankings improve over time."),
    ("Quality over quantity", "One well-written article on a high-DA platform is worth more than 10 low-quality posts."),
    ("Diversify your strategy", "Blog posting is one part of SEO. Also focus on Google Business Profile, social media, and on-page SEO."),
    ("Keep content fresh", "Update your keywords and templates regularly to stay relevant to current trends."),
]

for i, (title, detail) in enumerate(tips, 1):
    p = doc.add_paragraph()
    run = p.add_run(f"{i}. {title}: ")
    run.bold = True
    p.add_run(detail)

# ============================================================
# SAVE
# ============================================================
output_path = os.path.join("blog_poster", "output", "Blog_Auto_Poster_Documentation.docx")
os.makedirs(os.path.dirname(output_path), exist_ok=True)
doc.save(output_path)
print(f"Document saved: {output_path}")

# Open it
import platform as plat
import subprocess
filepath = os.path.abspath(output_path)
print(f"Opening: {filepath}")
try:
    if plat.system() == "Windows":
        os.startfile(filepath)
except Exception as e:
    print(f"Open manually: {filepath}")
